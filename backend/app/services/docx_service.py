import os
import uuid
import logging
import re
import json
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.ext.asyncio import AsyncSession
from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.models.pv import PV
from app.models.meeting import Meeting
from app.models.meeting_room import MeetingRoom
from app.models.action import Action
from app.models.setting import BrandingSettings

logger = logging.getLogger(__name__)

class DOCXService:
    def __init__(self, db: AsyncSession):
        self.db = db

    def _set_rtl(self, p):
        """
        Light RTL: Sets right alignment and professional Arabic font.
        Keeps logical LTR flow to prevent 'backwards text' issues in OnlyOffice.
        """
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for run in p.runs:
            # FreeSerif is installed and looks great for Arabic
            run.font.name = 'FreeSerif'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            
            # Set complex script font and basic font for compatibility
            rFonts.set(qn('w:cs'), 'FreeSerif')
            rFonts.set(qn('w:ascii'), 'FreeSerif')
            rFonts.set(qn('w:hAnsi'), 'FreeSerif')

            # Ensure font size is consistent (24 half-points = 12pt)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '24')
                rPr.append(sz)
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '24')
                rPr.append(szCs)

    def _set_document_rtl(self, doc):
        """
        Light Document RTL: Sets default fonts and some document-wide styles.
        """
        if 'Normal' in doc.styles:
            style = doc.styles['Normal']
            style.font.name = 'FreeSerif'
            style.font.size = Pt(12)

    async def generate_pv_docx(self, pv_id: str, client_id: str, branding_id: Optional[str] = None, language: str = "fr") -> str:
        """
        Generates a PV as a DOCX file and returns the file path.
        """
        # 0. Localization Strings
        LOCALES = {
            "ar": {
                "title": "محضر اجتماع",
                "date": "التاريخ",
                "location": "المكان",
                "duration": "المدة",
                "minutes": "دقيقة",
                "participants": "المشاركون",
                "agenda": "جدول الأعمال",
                "discussion": "ملخص المناقشات",
                "decisions": "القرارات",
                "actions": "خطة العمل",
                "task": "المهمة",
                "assignee": "المسؤول",
                "due_date": "الموعد النهائي",
                "signature": "الاعتماد",
                "director": "المدير العام",
                "not_assigned": "غير محدد",
            },
            "fr": {
                "title": "Procès-Verbal",
                "date": "Date",
                "location": "Lieu",
                "duration": "Durée",
                "minutes": "min",
                "participants": "Participants",
                "agenda": "Ordre du Jour",
                "discussion": "Résumé des Discussions",
                "decisions": "Décisions",
                "actions": "Plan d'Action",
                "task": "Tâche",
                "assignee": "Responsable",
                "due_date": "Échéance",
                "signature": "Approbation (Signature)",
                "director": "Directeur Général (DG)",
                "not_assigned": "Non défini",
            },
            "en": {
                "title": "Meeting Minutes",
                "date": "Date",
                "location": "Location",
                "duration": "Duration",
                "minutes": "min",
                "participants": "Participants",
                "agenda": "Agenda",
                "discussion": "Discussion Summary",
                "decisions": "Decisions",
                "actions": "Action Items",
                "task": "Task",
                "assignee": "Assignee",
                "due_date": "Due Date",
                "signature": "Approval (Signature)",
                "director": "General Manager (DG)",
                "not_assigned": "N/A",
            }
        }
        
        strings = LOCALES.get(language, LOCALES["fr"])
        is_rtl = (language == "ar")

        # 1. Load Data
        stmt = (
            select(PV)
            .options(
                selectinload(PV.meeting).selectinload(Meeting.participants),
                selectinload(PV.meeting).selectinload(Meeting.agendas),
                selectinload(PV.meeting).selectinload(Meeting.room),
                selectinload(PV.sections),
            )
            .where(PV.id == pv_id)
            .where(PV.client_id == client_id)
        )
        result = await self.db.execute(stmt)
        pv_obj = result.scalar_one_or_none()

        if not pv_obj:
            raise HTTPException(status_code=404, detail="PV not found")

        from app.models.action import Assignment
        action_stmt = (
            select(Action)
            .options(selectinload(Action.assignments).selectinload(Assignment.user))
            .where(Action.meeting_id == pv_obj.meeting_id)
            .where(Action.client_id == client_id)
        )
        action_result = await self.db.execute(action_stmt)
        actions = action_result.scalars().all()

        # 2. Translation Logic
        display_title = pv_obj.title
        display_discussion = pv_obj.content_html
        display_decisions = [s.content for s in pv_obj.sections if s.type == "decision"]
        display_actions = []
        for a in actions:
            assignee_name = strings["not_assigned"]
            if a.assignments and len(a.assignments) > 0:
                assignment = a.assignments[0]
                if assignment.user and assignment.user.full_name:
                    assignee_name = assignment.user.full_name
                elif assignment.external_name:
                    assignee_name = assignment.external_name
                    
            display_actions.append({
                "description": a.title,
                "assignee": assignee_name,
                "due_date": a.due_date.strftime("%Y-%m-%d") if a.due_date else datetime.now().strftime("%Y-%m-%d")
            })

        from app.services.pv_service import PVService, TranslationError
        if pv_obj.language != language:
            logger.info(f"Translating DOCX content from {pv_obj.language} to {language} via Mistral...")
            content_to_translate = {
                "title": display_title,
                "summary": display_discussion,
                "decisions": display_decisions,
                "actions": display_actions
            }
            try:
                translated = await PVService.translate_content(content_to_translate, language)
                display_title = translated.get("title", display_title)
                display_discussion = translated.get("summary", display_discussion)
                display_decisions = translated.get("decisions", display_decisions)
                display_actions = translated.get("actions", display_actions)
            except TranslationError as e:
                logger.error(f"Mistral translation failed for DOCX: {e}. Falling back to original.")

        # 3. Create Document
        doc = Document()
        
        # Title
        title_para = doc.add_heading(strings["title"], 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_pv_title = doc.add_paragraph(display_title)
        p_pv_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p_pv_title.runs:
            p_pv_title.runs[0].bold = True

        sep = doc.add_paragraph("_" * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta info
        start_time = pv_obj.meeting.start_time
        end_time = pv_obj.meeting.end_time
        date_str = start_time.strftime("%Y-%m-%d") if start_time else "N/A"
        duration_val = "N/A"
        
        if start_time and end_time:
            diff = end_time - start_time
            duration_val = str(max(1, int(diff.total_seconds() / 60)))

        meta = doc.add_paragraph()
        meta.add_run(f"{strings['date']}: ").bold = True
        meta.add_run(f"{date_str}    ")
        meta.add_run(f"{strings['duration']}: ").bold = True
        meta.add_run(f"{duration_val} {strings['minutes']}    ")
        meta.add_run(f"{strings['location']}: ").bold = True
        
        display_location = "N/A"
        if pv_obj.meeting.room and pv_obj.meeting.room.name:
            display_location = pv_obj.meeting.room.name
        elif pv_obj.meeting.location:
            display_location = pv_obj.meeting.location
            
        meta.add_run(f"{display_location}")

        # Participants
        p_part = doc.add_paragraph()
        p_part.add_run(f"{strings['participants']}: ").bold = True
        p_part.add_run(", ".join([p.name or p.email for p in pv_obj.meeting.participants]))

        # Agenda
        h_agenda = doc.add_heading(strings["agenda"], level=1)
        for a in sorted(pv_obj.meeting.agendas, key=lambda x: x.order):
            doc.add_paragraph(a.title, style='List Bullet')

        # Discussion
        h_disc = doc.add_heading(strings["discussion"], level=1)
        
        clean_text = re.sub('<[^<]+?>', '', display_discussion or "")
        for line in clean_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        # Decisions
        if display_decisions:
            h_dec = doc.add_heading(strings["decisions"], level=1)
            for d in display_decisions:
                doc.add_paragraph(d, style='List Bullet')

        # Actions
        if display_actions:
            h_act = doc.add_heading(strings["actions"], level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = strings["task"]
            hdr_cells[1].text = strings["assignee"]
            hdr_cells[2].text = strings["due_date"]

            for action in display_actions:
                row_cells = table.add_row().cells
                row_cells[0].text = str(action.get("description", "N/A"))
                row_cells[1].text = str(action.get("assignee", "N/A"))
                row_cells[2].text = str(action.get("due_date", "N/A"))

        # Signature
        doc.add_paragraph("\n" * 3)
        doc.add_paragraph(f"{strings['signature']}:")
        doc.add_paragraph("___________________________")
        doc.add_paragraph(strings["director"])

        # Metadata
        if duration_val != "N/A":
            doc.core_properties.comments = f"Meeting Duration: {duration_val} minutes"

        # FINAL LIGHT PASS: Apply Alignment and Font for Arabic
        if is_rtl:
            # 1. Global Document adjustments
            self._set_document_rtl(doc)
            
            # 2. All Paragraphs
            for p in doc.paragraphs:
                self._set_rtl(p)
                
            # 3. All Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._set_rtl(p)

        # Save
        filename = f"pv_{pv_id}_{uuid.uuid4().hex[:8]}.docx"
        filepath = f"/tmp/{filename}"
        doc.save(filepath)
        
        return filepath
