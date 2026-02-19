import logging
from typing import List, Optional, Dict, Any
from datetime import datetime

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import Session
from fastapi import HTTPException, status
from docx import Document
from fpdf import FPDF

from backend.app.core.config import settings
from backend.app.models.transcription import Transcription
from backend.app.models.recording import Recording
from backend.app.models.meeting import Meeting
from backend.app.models.user import User, UserRole # Import User and UserRole
from backend.app.schemas.transcription import (
    TranscriptionCreate, TranscriptionUpdate, TranscriptionStatus,
    SpeakerSegment, WordTimestamp
)
from backend.app.services.whisper_client import whisper_client
from backend.app.utils.storage import storage_service
import logging
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone
import datetime as dt # Alias datetime for timedelta usage

logger = logging.getLogger(__name__)

class TranscriptionService:
    async def start_transcription(
        self,
        db: AsyncSession,
        recording_id: int,
        current_user_id: int,
        language: Optional[str] = None,
        enable_diarization: bool = False
    ) -> Transcription:
        """
        Starts a transcription job for a given recording.
        """
        result = await db.execute(select(Recording).where(Recording.id == recording_id))
        recording = result.scalar_one_or_none()
        if not recording:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Recording not found")

        # Check if the user has access to the meeting associated with the recording
        result = await db.execute(select(Meeting).where(Meeting.id == recording.meeting_id))
        meeting = result.scalar_one_or_none()
        if not meeting or not await self.user_can_access_meeting(db, current_user_id, meeting):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not authorized to access this recording's meeting")

        # Ensure the recording has a file_path
        if not recording.file_path:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Recording has no associated file for transcription")

        # Create a new transcription entry in PENDING status
        db_transcription = Transcription(
            meeting_id=meeting.id,  # Add meeting_id
            recording_id=recording_id,
            language=language,
            status=TranscriptionStatus.PENDING,
            created_by_id=current_user_id,
            started_at=datetime.now(timezone.utc)
        )
        try:
            db.add(db_transcription)
            await db.commit()
            await db.refresh(db_transcription)

            # Asynchronously call the Whisper API
            try:
                # Get a signed URL for the audio file
                audio_file_url = await storage_service.get_s3_download_url(recording.file_path)
                
                whisper_result = await whisper_client.call_whisper_api(
                    audio_file_url=audio_file_url,
                    language=language,
                    enable_diarization=enable_diarization
                )
                
                # Process the result
                await self.process_transcription_result(db, db_transcription.id, whisper_result)
                
            except Exception as e:
                logger.error(f"Failed to start transcription for recording {recording_id}: {e}")
                db_transcription.status = TranscriptionStatus.FAILED
                db_transcription.failed_reason = str(e)
                db_transcription.completed_at = datetime.now(timezone.utc)
                db.add(db_transcription)
                await db.commit()
                await db.refresh(db_transcription)
                raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Transcription failed to start: {e}")

            return db_transcription
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to create transcription entry for recording {recording_id}: {e}")
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to create transcription entry: {e}")

    async def get_transcription_by_id(self, db: AsyncSession, transcription_id: int, current_user_id: int) -> Transcription:
        """
        Retrieves a transcription by its ID.
        """
        result = await db.execute(select(Transcription).where(Transcription.id == transcription_id))
        transcription = result.scalar_one_or_none()
        if not transcription:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcription not found")
        
        # Check user authorization
        result = await db.execute(select(Recording).where(Recording.id == transcription.recording_id))
        recording = result.scalar_one_or_none()
        result = await db.execute(select(Meeting).where(Meeting.id == recording.meeting_id))
        meeting = result.scalar_one_or_none()
        if not meeting or not await self.user_can_access_meeting(db, current_user_id, meeting):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not authorized to access this transcription")

        return transcription

    async def get_transcriptions_by_meeting(self, db: AsyncSession, meeting_id: int, current_user_id: int) -> List[Transcription]:
        """
        Retrieves all transcriptions for a given meeting.
        """
        result = await db.execute(select(Meeting).where(Meeting.id == meeting_id))
        meeting = result.scalar_one_or_none()
        if not meeting or not await self.user_can_access_meeting(db, current_user_id, meeting):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not authorized to access this meeting's transcriptions")

        result = await db.execute(select(Recording).where(Recording.meeting_id == meeting_id))
        recordings = result.scalars().all()
        recording_ids = [rec.id for rec in recordings]
        
        result = await db.execute(select(Transcription).where(Transcription.recording_id.in_(recording_ids)))
        transcriptions = result.scalars().all()
        return transcriptions

    async def update_transcription(
        self,
        db: AsyncSession,
        transcription_id: int,
        transcription_update: TranscriptionUpdate,
        current_user_id: int
    ) -> Transcription:
        """
        Updates an existing transcription.
        """
        db_transcription = await self.get_transcription_by_id(db, transcription_id, current_user_id) # Includes auth check

        update_data = transcription_update.dict(exclude_unset=True)
        for key, value in update_data.items():
            if key == "speaker_segments" and value is not None:
                db_transcription.speaker_segments = [s.dict() for s in value]
            elif key == "word_timestamps" and value is not None:
                db_transcription.word_timestamps = [w.dict() for w in value]
            else:
                setattr(db_transcription, key, value)
        
        db_transcription.updated_at = datetime.now(timezone.utc)
        if db_transcription.status != TranscriptionStatus.FAILED:
            db_transcription.status = TranscriptionStatus.EDITED # Mark as edited if manually updated

        try:
            db.add(db_transcription)
            await db.commit()
            await db.refresh(db_transcription)
            return db_transcription
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to update transcription {transcription_id}: {e}")
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to update transcription: {e}")

    async def delete_transcription(self, db: AsyncSession, transcription_id: int, current_user_id: int):
        """
        Deletes a transcription.
        """
        db_transcription = await self.get_transcription_by_id(db, transcription_id, current_user_id) # Includes auth check
        try:
            await db.delete(db_transcription)
            await db.commit()
            return {"message": "Transcription deleted successfully"}
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to delete transcription {transcription_id}: {e}")
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to delete transcription: {e}")

    async def process_transcription_result(self, db: AsyncSession, transcription_id: int, whisper_result: Dict[str, Any]):
        """
        Processes the result from the Whisper API and updates the transcription in the database.
        """
        result = await db.execute(select(Transcription).where(Transcription.id == transcription_id))
        db_transcription = result.scalar_one_or_none()
        if not db_transcription:
            logger.error(f"Transcription with ID {transcription_id} not found for processing result.")
            return

        db_transcription.transcribed_text = whisper_result.get("text")
        db_transcription.language = whisper_result.get("language", db_transcription.language)
        # db_transcription.duration = whisper_result.get("duration", db_transcription.duration) # Transcription model has no duration field
        
        if "segments" in whisper_result:
            db_transcription.speaker_segments = [
                SpeakerSegment(
                    speaker=s.get("speaker"),
                    text=s["text"],
                    start=s["start"],
                    end=s["end"]
                ).dict() for s in whisper_result["segments"]
            ]
        
        if "word_timestamps" in whisper_result:
            db_transcription.word_timestamps = [
                WordTimestamp(
                    word=w["word"],
                    start=w["start"],
                    end=w["end"],
                    confidence=w.get("confidence")
                ).dict() for w in whisper_result["word_timestamps"]
            ]

        db_transcription.status = TranscriptionStatus.COMPLETED
        db_transcription.completed_at = datetime.now(timezone.utc)
        try:
            db.add(db_transcription)
            await db.commit()
            await db.refresh(db_transcription)
            logger.info(f"Transcription {transcription_id} processed and updated successfully.")
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to process transcription result for {transcription_id}: {e}")
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to process transcription result: {e}")

    async def detect_language(self, audio_file_url: str) -> Optional[str]:
        """
        Detects the language of an audio file using the Whisper API.
        This might be a separate endpoint on the Whisper service or part of the transcription.
        For now, we'll assume `call_whisper_api` can return it.
        """
        try:
            # Call Whisper API with a minimal request to get language detection
            # Assuming the Whisper API can return language without full transcription if needed
            # Or, we can just rely on the language returned by a full transcription
            result = await whisper_client.call_whisper_api(audio_file_url=audio_file_url, language=None, enable_diarization=False)
            return result.get("language")
        except Exception as e:
            logger.warning(f"Could not detect language for {audio_file_url}: {e}")
            return None

    async def user_can_access_meeting(self, db: AsyncSession, user_id: int, meeting: Meeting) -> bool:
        """
        Checks if a user has access to a specific meeting.
        A user has access if they are the organizer or an admin.
        """
        if meeting.organizer_id == user_id:
            return True
        
        # Check if user is an admin
        result = await db.execute(select(User).where(User.id == user_id))
        user = result.scalar_one_or_none()
        
        if user and user.role == UserRole.ADMIN:
            return True
            
        return False

    def format_transcription(self, transcription: Transcription, format_type: str = "text") -> str:
        """
        Formats the transcription into different output formats.
        """
        if not transcription.transcribed_text:
            return "No transcribed text available."

        if format_type == "text" or format_type == "txt":
            return transcription.transcribed_text
        elif format_type == "srt":
            return self._format_to_srt(transcription)
        elif format_type == "vtt":
            return self._format_to_vtt(transcription)
        elif format_type == "json":
            # Assuming transcription is a Pydantic model or has a .json() method
            # If it's a SQLAlchemy model, you might need to convert it to a Pydantic schema first
            return transcription.json(indent=2) 
        else:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Unsupported format type: {format_type}")

    def _format_to_srt(self, transcription: Transcription) -> str:
        """Helper to format transcription to SRT."""
        if not transcription.speaker_segments:
            return "SRT format requires speaker segments or word timestamps."

        srt_content = []
        for i, segment_data in enumerate(transcription.speaker_segments):
            segment = SpeakerSegment(**segment_data)
            start_time = str(dt.timedelta(seconds=segment.start))
            end_time = str(dt.timedelta(seconds=segment.end))
            
            # SRT time format: HH:MM:SS,ms
            start_srt = start_time.split('.')[0] + ',' + start_time.split('.')[1][:3].ljust(3, '0') if '.' in start_time else start_time + ',000'
            end_srt = end_time.split('.')[0] + ',' + end_time.split('.')[1][:3].ljust(3, '0') if '.' in end_time else end_time + ',000'

            srt_content.append(str(i + 1))
            srt_content.append(f"{start_srt} --> {end_srt}")
            speaker_prefix = f"{segment.speaker}: " if segment.speaker else ""
            srt_content.append(f"{speaker_prefix}{segment.text}")
            srt_content.append("")
        return "\n".join(srt_content)

    def _format_to_vtt(self, transcription: Transcription) -> str:
        """Helper to format transcription to VTT."""
        if not transcription.speaker_segments:
            return "VTT format requires speaker segments or word timestamps."

        vtt_content = ["WEBVTT\n"]
        for i, segment_data in enumerate(transcription.speaker_segments):
            segment = SpeakerSegment(**segment_data)
            start_time = str(dt.timedelta(seconds=segment.start))
            end_time = str(dt.timedelta(seconds=segment.end))

            # VTT time format: HH:MM:SS.ms
            start_vtt = start_time.split('.')[0] + '.' + start_time.split('.')[1][:3].ljust(3, '0') if '.' in start_time else start_time + '.000'
            end_vtt = end_time.split('.')[0] + '.' + end_time.split('.')[1][:3].ljust(3, '0') if '.' in end_time else end_time + '.000'

            vtt_content.append(f"{start_vtt} --> {end_vtt}")
            speaker_prefix = f"<v {segment.speaker}> " if segment.speaker else ""
            vtt_content.append(f"{speaker_prefix}{segment.text}")
            vtt_content.append("")
        return "\n".join(vtt_content)

    def export_transcription(self, transcription: Transcription, export_format: str) -> bytes:
        """
        Exports the transcription into various file formats (TXT, DOCX, PDF).
        Returns the file content as bytes.
        """
        if not transcription.transcribed_text:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No transcribed text to export.")

        if export_format == "txt":
            return transcription.transcribed_text.encode('utf-8')
        elif export_format == "docx":
            document = Document()
            document.add_heading('Transcription', level=1)
            if transcription.speaker_segments:
                for segment_data in transcription.speaker_segments:
                    segment = SpeakerSegment(**segment_data)
                    speaker_text = f"{segment.speaker}: " if segment.speaker else ""
                    document.add_paragraph(f"{speaker_text}{segment.text}")
            else:
                document.add_paragraph(transcription.transcribed_text)
            
            # Save document to a bytes buffer
            from io import BytesIO
            byte_io = BytesIO()
            document.save(byte_io)
            byte_io.seek(0)
            return byte_io.getvalue()
        elif export_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # FPDF does not directly support complex text layout or UTF-8 out of the box for all fonts.
            # Need to ensure a font that supports UTF-8 is used.
            # For simplicity, using a basic font and handling potential encoding issues.
            # A more robust solution might involve a library like ReportLab or a custom font.
            
            # Add a font that supports UTF-8 characters (e.g., DejaVuSansCondensed)
            # This requires the font file to be available. For a simple example, we might skip complex chars.
            # For production, you'd need to manage font files.
            try:
                pdf.add_font('DejaVuSansCondensed', '', 'DejaVuSansCondensed.ttf', uni=True)
                pdf.set_font('DejaVuSansCondensed', size=12)
            except Exception:
                logger.warning("DejaVuSansCondensed font not found, using default Arial. UTF-8 characters might not display correctly in PDF.")
                pdf.set_font("Arial", size=12) # Fallback to Arial if custom font not available

            pdf.multi_cell(0, 10, "Transcription")
            pdf.ln(10)
            
            if transcription.speaker_segments:
                for segment_data in transcription.speaker_segments:
                    segment = SpeakerSegment(**segment_data)
                    speaker_text = f"{segment.speaker}: " if segment.speaker else ""
                    pdf.multi_cell(0, 10, f"{speaker_text}{segment.text}")
            else:
                pdf.multi_cell(0, 10, transcription.transcribed_text)
            
            return pdf.output(dest='S').encode('latin-1') # FPDF outputs bytes, encode as latin-1 for common PDF viewers
        else:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Unsupported export format: {export_format}")

transcription_service = TranscriptionService()
