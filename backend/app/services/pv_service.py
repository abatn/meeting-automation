import logging
from datetime import datetime, timezone
from typing import Optional
from docx import Document
from docx.shared import Inches
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from backend.app.core.config import settings
from backend.app.models.meeting import Meeting
from backend.app.models.pv import PV, PVStatus
from backend.app.models.transcription import Transcription
from backend.app.models.user import User, UserRole
from backend.app.schemas.pv import PVCreate, PVUpdate
from backend.app.services.mistral_client import MistralClient, MOCK_PV_RESPONSE
from backend.app.utils.storage import storage_service
from sqlalchemy.orm import selectinload, Session
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO

logger = logging.getLogger(__name__)

mistral_client = MistralClient()

class PVService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def generate_pv(self, meeting_id: int, transcription_id: Optional[int], template: Optional[str], current_user: User) -> PV:
        result = await self.db.execute(select(Meeting).filter(Meeting.id == meeting_id))
        meeting = result.scalar_one_or_none()
        if not meeting:
            raise HTTPException(status_code=404, detail="Meeting not found")

        if transcription_id:
            result = await self.db.execute(select(Transcription).filter(Transcription.id == transcription_id))
            transcription = result.scalar_one_or_none()
            if not transcription:
                raise HTTPException(status_code=404, detail="Transcription not found")
        else:
            result = await self.db.execute(select(Transcription).filter(Transcription.meeting_id == meeting_id))
            transcription = result.scalar_one_or_none()
            if not transcription:
                raise HTTPException(status_code=400, detail="No transcription found for this meeting.")

        if not transcription.transcribed_text:
            raise HTTPException(status_code=400, detail="Transcription content is empty.")

        try:
            pv_content = await mistral_client.generate_pv(transcription.transcribed_text)
            decisions = await mistral_client.extract_decisions(pv_content)
            action_points = await mistral_client.extract_action_items(pv_content)
        except Exception as e:
            logger.error(f"Mistral API error: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to generate PV: {e}")

        pv = PV(
            title=f"PV for Meeting '{meeting.title}' on {meeting.date.strftime('%Y-%m-%d')}",
            content=pv_content,
            meeting_id=meeting_id,
            generated_by_id=current_user.id,
            decisions=decisions,
            action_points=action_points,
            status=PVStatus.DRAFT
        )
        self.db.add(pv)
        await self.db.commit()
        await self.db.refresh(pv)
        
        return pv

    async def get_pv_by_id(self, pv_id: int, current_user: User) -> Optional[PV]:
        logger.info(f"Fetching PV with ID: {pv_id}")
        result = await self.db.execute(
            select(PV)
            .filter(PV.id == pv_id)
        )
        pv = result.scalar_one_or_none()
        if not pv:
            logger.warning(f"PV with ID: {pv_id} not found.")
            raise HTTPException(status_code=404, detail="PV not found")
        logger.info(f"PV with ID: {pv_id} found.")
        return pv

    async def get_pvs_by_meeting(self, meeting_id: int, current_user: User) -> list[PV]:
        logger.info(f"Fetching all PVs for meeting ID: {meeting_id}")
        result = await self.db.execute(
            select(PV)
            .filter(PV.meeting_id == meeting_id)
        )
        pvs = result.scalars().all()
        logger.info(f"Found {len(pvs)} PVs for meeting ID: {meeting_id}")
        return pvs

    async def update_pv(self, pv_id: int, pv_update: PVUpdate, current_user: User) -> PV:
        pv = await self.get_pv_by_id(pv_id, current_user)
        update_data = pv_update.dict(exclude_unset=True)
        for key, value in update_data.items():
            setattr(pv, key, value)
        await self.db.commit()
        await self.db.refresh(pv)
        return pv

    async def validate_pv(self, pv_id: int, user: User, comment: Optional[str]) -> PV:
        user = await self.db.merge(user)
        query = select(PV).where(PV.id == pv_id)
        result = await self.db.execute(query)
        pv = result.scalar_one_or_none()
        
        if not pv:
            raise HTTPException(status_code=404, detail="PV not found")

        if user.role != UserRole.DG:
            raise HTTPException(status_code=403, detail="Only DGs can validate PVs")

        if pv.status == PVStatus.VALIDATED:
            raise HTTPException(status_code=400, detail="PV is already validated.")

        pv.validated_at = datetime.now(timezone.utc)
        pv.validator_id = user.id
        pv.validation_comment = comment
        pv.status = PVStatus.VALIDATED
        self.db.add(pv)
        await self.db.commit()
        await self.db.refresh(pv)
        
        # Eagerly load the validator relationship
        result = await self.db.execute(
            select(PV).options(selectinload(PV.validator)).filter(PV.id == pv_id)
        )
        return result.scalar_one()

    async def delete_pv(self, pv_id: int, current_user: User):
        if current_user.role != UserRole.ADMIN:
            raise HTTPException(status_code=403, detail="Only Admins can delete PVs")
        pv = await self.get_pv_by_id(pv_id, current_user)
        await self.db.delete(pv)
        await self.db.commit()
        return {"message": "PV deleted successfully"}

    async def generate_pv_pdf(self, pv_id: int, current_user: User):
        pv = await self.get_pv_by_id(pv_id, current_user)
        result = await self.db.execute(select(Meeting).filter(Meeting.id == pv.meeting_id))
        meeting = result.scalars().first()
        if not meeting:
            raise HTTPException(status_code=404, detail="Meeting not found for this PV")

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(f"<h1>Procès-verbal</h1>", styles['h1']))
        story.append(Paragraph(f"<h2>{meeting.title}</h2>", styles['h2']))
        story.append(Paragraph(pv.content, styles['Normal']))

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=pv_{pv_id}.pdf"})

    async def generate_pv_docx(self, pv_id: int, current_user: User):
        pv = await self.get_pv_by_id(pv_id, current_user)
        result = await self.db.execute(select(Meeting).filter(Meeting.id == pv.meeting_id))
        meeting = result.scalars().first()
        if not meeting:
            raise HTTPException(status_code=404, detail="Meeting not found for this PV")

        document = Document()
        document.add_heading('Procès-verbal', level=1)
        document.add_heading(meeting.title, level=2)
        document.add_paragraph(pv.content)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=pv_{pv_id}.docx"})

async def extract_decisions(pv_content: str) -> str:
    try:
        decisions_content = await mistral_client.extract_decisions(pv_content)
        if not decisions_content:
            logger.error("Mistral API returned empty content for decision extraction.")
            raise HTTPException(status_code=500, detail="Mistral API returned empty content for decision extraction.")
        return decisions_content
    except Exception as e:
        logger.error(f"Error extracting decisions from PV: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to extract decisions: {e}. Details: {e}")

async def extract_action_points(pv_content: str) -> str:
    try:
        action_points_content = await mistral_client.extract_action_items(pv_content)
        if not action_points_content:
            raise HTTPException(status_code=500, detail="Mistral API returned empty content for action points extraction.")
        return action_points_content
    except Exception as e:
        logger.error(f"Error extracting action points from PV: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract action points: {e}")

async def generate_summary_from_pv(pv_content: str) -> str:
    try:
        summary_content = await mistral_client.generate_summary(pv_content)
        if not summary_content:
            raise HTTPException(status_code=500, detail="Mistral API returned empty content for summary generation.")
        return summary_content
    except Exception as e:
        logger.error(f"Error generating summary from PV: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate summary: {e}")